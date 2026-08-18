#!/usr/bin/env python3
"""Build the Chinese operator manual as a polished, deterministic DOCX.

Design contract:
- compact_reference_guide preset
- exactly one editorial_cover opening pattern
- US Letter, 1 inch margins, explicit list numbering and table geometry
"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
MARKDOWN_PATH = PROJECT_ROOT / "docs" / "离线SBOM证据工作台_用户使用手册.md"
DOCX_PATH = PROJECT_ROOT / "docs" / "离线SBOM证据工作台_用户使用手册.docx"

PRESET_NAME = "compact_reference_guide"
HEADER_PATTERN = "editorial_cover"
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
MUTED = "5D6B78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
GOLD = "96721C"
WHITE = "FFFFFF"
BLACK = "000000"
CODE_FILL = "F6F8FA"
CODE_BORDER = "D8DEE4"
RISK_RED = "9B1C1C"


def _set_font(run, *, name="Calibri", east_asia="PingFang SC", size=None,
              color=None, bold=None, italic=None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style, *, name="Calibri", east_asia="PingFang SC",
                    size=11, color=BLACK, bold=False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)


def _set_spacing(style, *, before=0, after=0, line=1.0,
                 keep_with_next=False) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_with_next
    fmt.widow_control = True


def _shade(element, fill: str) -> None:
    properties = element.get_or_add_pPr() if element.tag == qn("w:p") else element.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def _paragraph_border(paragraph, *, color=CODE_BORDER, size="6",
                      left=True, top=True, right=True, bottom=True) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    for edge, enabled in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        if not enabled:
            continue
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "3")
        node.set(qn("w:color"), color)


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    _set_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, value, end))
    tail = paragraph.add_run(" 页")
    _set_font(tail, size=9, color=MUTED)


def _set_cell_margins(cell) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    values = {
        "top": CELL_TOP_BOTTOM_DXA,
        "bottom": CELL_TOP_BOTTOM_DXA,
        "start": CELL_START_END_DXA,
        "end": CELL_START_END_DXA,
    }
    for name, value in values.items():
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError("table widths must total 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)
        for index, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[index]))
            tcw.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _repeat_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def _configure_numbering(document: Document) -> dict[str, int]:
    numbering = document.part.numbering_part.element
    existing_abstract = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_abstract = max(existing_abstract, default=0) + 1
    next_num = max(existing_num, default=0) + 1
    result: dict[str, int] = {}

    definitions = (
        ("bullet", "bullet", "•", "Symbol"),
        ("decimal", "decimal", "%1.", "Calibri"),
        ("check", "bullet", "☐", "Arial Unicode MS"),
    )
    for offset, (name, number_format, text, font) in enumerate(definitions):
        abstract_id = next_abstract + offset
        num_id = next_num + offset
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), number_format)
        level.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        level.append(lvl_text)
        justification = OxmlElement("w:lvlJc")
        justification.set(qn("w:val"), "left")
        level.append(justification)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "271")
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        level.append(ppr)
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font)
        fonts.set(qn("w:hAnsi"), font)
        rpr.append(fonts)
        level.append(rpr)
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)
        result[name] = num_id
    return result


def _apply_numbering(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    numpr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    identifier = OxmlElement("w:numId")
    identifier.set(qn("w:val"), str(num_id))
    numpr.extend((level, identifier))
    ppr.append(numpr)


def _new_number_instance(document: Document, base_num_id: int) -> int:
    numbering = document.part.numbering_part.element
    base = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == base_num_id
    )
    abstract_id = int(base.find(qn("w:abstractNumId")).get(qn("w:val")))
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _set_style_font(normal, size=11, color=BLACK)
    _set_spacing(normal, before=0, after=6, line=1.25)

    title = styles["Title"]
    _set_style_font(title, size=30, color=INK_BLUE, bold=True)
    _set_spacing(title, before=0, after=8, line=1.0, keep_with_next=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = styles["Subtitle"]
    _set_style_font(subtitle, size=15, color=DARK_BLUE)
    _set_spacing(subtitle, before=0, after=8, line=1.15, keep_with_next=True)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        _set_style_font(style, size=size, color=color, bold=True)
        _set_spacing(style, before=before, after=after, line=1.1, keep_with_next=True)
        style.paragraph_format.page_break_before = False

    for style_name in ("List Paragraph",):
        style = styles[style_name]
        _set_style_font(style, size=11, color=BLACK)
        _set_spacing(style, before=0, after=4, line=1.25)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", 1)
    else:
        code = styles["Code Block"]
    _set_style_font(code, name="Menlo", east_asia="PingFang SC", size=8.3, color=INK_BLUE)
    _set_spacing(code, before=0, after=0, line=1.05)
    code.paragraph_format.left_indent = Inches(0.08)
    code.paragraph_format.right_indent = Inches(0.08)
    code.paragraph_format.keep_together = True

    if "Callout" not in styles:
        callout = styles.add_style("Callout", 1)
    else:
        callout = styles["Callout"]
    _set_style_font(callout, size=10.5, color=INK_BLUE)
    _set_spacing(callout, before=6, after=8, line=1.2)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.12)

    if "Table Text" not in styles:
        table_text = styles.add_style("Table Text", 1)
    else:
        table_text = styles["Table Text"]
    _set_style_font(table_text, size=9.3, color=BLACK)
    _set_spacing(table_text, before=0, after=0, line=1.15)

    if "Table Header" not in styles:
        table_header = styles.add_style("Table Header", 1)
    else:
        table_header = styles["Table Header"]
    _set_style_font(table_header, size=9.3, color=INK_BLUE, bold=True)
    _set_spacing(table_header, before=0, after=0, line=1.15)


def _configure_sections(document: Document) -> None:
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run("离线 SBOM 证据工作台  |  用户使用手册  |  v0.4.0")
        _set_font(run, size=8.5, color=MUTED, bold=True)

        first_header = section.first_page_header
        first_header.paragraphs[0].text = ""

        footer = section.footer
        _add_page_field(footer.paragraphs[0])
        first_footer = section.first_page_footer
        first_footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        first_run = first_footer.paragraphs[0].add_run("SELF_TEST_NOT_CUSTOMER_EVIDENCE")
        _set_font(first_run, size=8.5, color=MUTED, bold=True)


def _add_inline_markup(paragraph, text: str, *, default_size=None,
                       default_color=BLACK) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            _set_font(run, size=default_size, color=default_color)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_font(
                run,
                name="Menlo",
                east_asia="PingFang SC",
                size=9 if default_size is None else min(default_size, 9),
                color=INK_BLUE,
            )
        else:
            run = paragraph.add_run(token[2:-2])
            _set_font(run, size=default_size, color=default_color, bold=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _set_font(run, size=default_size, color=default_color)


def _add_editorial_cover(document: Document) -> None:
    """Add exactly one editorial_cover opening pattern."""

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)

    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("操作手册  /  LOCAL ENGINEERING SELF-TEST")
    _set_font(run, size=10, color=GOLD, bold=True)

    title = document.add_paragraph(style="Title")
    run = title.add_run("离线 SBOM 证据工作台")
    _set_font(run, size=30, color=INK_BLUE, bold=True)

    subtitle = document.add_paragraph(style="Subtitle")
    run = subtitle.add_run("从 M3A 到 M6A 的打开、运行、复核与恢复指南")
    _set_font(run, size=15, color=DARK_BLUE)

    descriptor = document.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_before = Pt(14)
    descriptor.paragraph_format.space_after = Pt(64)
    run = descriptor.add_run("Offline SBOM Evidence Workbench v0.4.0")
    _set_font(run, size=10.5, color=MUTED)

    badge = document.add_paragraph()
    badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge.paragraph_format.left_indent = Inches(0.85)
    badge.paragraph_format.right_indent = Inches(0.85)
    badge.paragraph_format.space_after = Pt(14)
    _shade(badge._p, LIGHT_BLUE)
    _paragraph_border(badge, color="B9C9D8", size="5")
    run = badge.add_run("SELF_TEST_NOT_CUSTOMER_EVIDENCE  ·  OPEN CANDIDATE")
    _set_font(run, size=10, color=INK_BLUE, bold=True)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run("适用平台：Apple Silicon macOS  |  手册日期：2026-08-04")
    _set_font(run, size=10, color=MUTED)

    boundary = document.add_paragraph()
    boundary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = boundary.add_run("机械验证不等于组件完整、制造商批准、CRA 符合、CAB 结论或认证")
    _set_font(run, size=9.5, color=RISK_RED, bold=True)

    page_break = document.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)


def _add_callout(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Callout")
    _shade(paragraph._p, CALLOUT_FILL)
    _paragraph_border(paragraph, color=BLUE, size="8", left=True, top=False, right=False, bottom=False)
    _add_inline_markup(paragraph, text, default_size=10.5, default_color=INK_BLUE)


def _add_code_block(document: Document, lines: list[str]) -> None:
    for index, line in enumerate(lines):
        paragraph = document.add_paragraph(style="Code Block")
        _shade(paragraph._p, CODE_FILL)
        _paragraph_border(
            paragraph,
            color=CODE_BORDER,
            size="4",
            top=index == 0,
            bottom=index == len(lines) - 1,
            left=True,
            right=True,
        )
        paragraph.paragraph_format.space_before = Pt(3 if index == 0 else 0)
        paragraph.paragraph_format.space_after = Pt(5 if index == len(lines) - 1 else 0)
        paragraph.paragraph_format.keep_with_next = index < len(lines) - 1
        run = paragraph.add_run(line if line else " ")
        _set_font(run, name="Menlo", east_asia="PingFang SC", size=8.3, color=INK_BLUE)


def _table_widths(column_count: int) -> list[int]:
    if column_count == 4:
        return [1680, 1840, 2500, 3340]
    if column_count == 3:
        return [3300, 3030, 3030]
    if column_count == 2:
        return [2700, 6660]
    base = CONTENT_WIDTH_DXA // column_count
    widths = [base] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows or not rows[0]:
        return
    column_count = len(rows[0])
    if any(len(row) != column_count for row in rows):
        raise ValueError("Markdown table has inconsistent column count")
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = "Table Header" if row_index == 0 else "Table Text"
            if column_index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline_markup(
                paragraph,
                value,
                default_size=9.3,
                default_color=INK_BLUE if row_index == 0 else BLACK,
            )
            if row_index == 0:
                _shade(cell._tc, LIGHT_BLUE)
    _repeat_header(table.rows[0])
    _set_table_geometry(table, _table_widths(column_count))
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    cursor = start
    while cursor < len(lines) and lines[cursor].lstrip().startswith("|"):
        cells = [cell.strip() for cell in lines[cursor].strip().strip("|").split("|")]
        if cursor == start + 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            cursor += 1
            continue
        rows.append(cells)
        cursor += 1
    return rows, cursor


def _add_markdown_body(document: Document, markdown: str, numbering: dict[str, int]) -> None:
    lines = markdown.splitlines()
    cursor = 0
    skipped_title = False
    in_code = False
    code_lines: list[str] = []
    code_fence = ""
    active_list_kind: str | None = None
    active_decimal_num_id: int | None = None

    while cursor < len(lines):
        raw = lines[cursor]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if not in_code:
                active_list_kind = None
                in_code = True
                code_fence = stripped[3:].strip()
                code_lines = []
            else:
                _add_code_block(document, code_lines)
                in_code = False
                code_fence = ""
                code_lines = []
            cursor += 1
            continue
        if in_code:
            code_lines.append(raw)
            cursor += 1
            continue

        if not stripped:
            cursor += 1
            continue

        if stripped.startswith("# "):
            active_list_kind = None
            if not skipped_title:
                skipped_title = True
            cursor += 1
            continue

        if stripped.startswith("> "):
            active_list_kind = None
            parts: list[str] = []
            while cursor < len(lines) and lines[cursor].strip().startswith("> "):
                parts.append(lines[cursor].strip()[2:].rstrip("  "))
                cursor += 1
            text = " ".join(parts)
            if any(prefix in text for prefix in ("适用对象：", "适用版本：", "参考对象：", "手册日期：")):
                continue
            _add_callout(document, text)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            active_list_kind = None
            level = len(heading_match.group(1)) - 1
            paragraph = document.add_paragraph(style=f"Heading {level}")
            _add_inline_markup(paragraph, heading_match.group(2), default_color=BLUE if level < 3 else DARK_BLUE)
            cursor += 1
            continue

        if stripped.startswith("|") and cursor + 1 < len(lines) and lines[cursor + 1].lstrip().startswith("|"):
            active_list_kind = None
            rows, cursor = _parse_table(lines, cursor)
            _add_table(document, rows)
            continue

        checklist = re.match(r"^- \[ \]\s+(.+)$", stripped)
        if checklist:
            active_list_kind = "check"
            paragraph = document.add_paragraph(style="List Paragraph")
            _apply_numbering(paragraph, numbering["check"])
            _add_inline_markup(paragraph, checklist.group(1))
            cursor += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            active_list_kind = "bullet"
            paragraph = document.add_paragraph(style="List Paragraph")
            _apply_numbering(paragraph, numbering["bullet"])
            _add_inline_markup(paragraph, bullet.group(1))
            cursor += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            if active_list_kind != "decimal" or active_decimal_num_id is None:
                active_decimal_num_id = _new_number_instance(document, numbering["decimal"])
            active_list_kind = "decimal"
            paragraph = document.add_paragraph(style="List Paragraph")
            _apply_numbering(paragraph, active_decimal_num_id)
            _add_inline_markup(paragraph, numbered.group(1))
            cursor += 1
            continue

        active_list_kind = None
        paragraph_parts = [stripped]
        cursor += 1
        while cursor < len(lines):
            upcoming = lines[cursor].strip()
            if (
                not upcoming
                or upcoming.startswith(("#", "- ", "> ", "```", "|"))
                or re.match(r"^\d+\.\s+", upcoming)
            ):
                break
            paragraph_parts.append(upcoming)
            cursor += 1
        paragraph = document.add_paragraph(style="Normal")
        _add_inline_markup(paragraph, " ".join(paragraph_parts))

    if in_code:
        raise ValueError(f"unclosed Markdown code fence: {code_fence}")


def _set_document_metadata(document: Document) -> None:
    properties = document.core_properties
    properties.title = "离线 SBOM 证据工作台用户使用手册"
    properties.subject = "M3A 至 M6A 本地工程自测操作指南"
    properties.author = "Offline SBOM Evidence Workbench Project"
    properties.keywords = "SBOM, CycloneDX, SPDX, EUVD, oMLX, evidence"
    properties.comments = "SELF_TEST_NOT_CUSTOMER_EVIDENCE"
    fixed_time = datetime(2026, 8, 4, 0, 0, 0, tzinfo=timezone.utc)
    properties.created = fixed_time
    properties.modified = fixed_time


def _audit_document(document: Document) -> None:
    section = document.sections[0]
    assert section.page_width.twips == PAGE_WIDTH_DXA
    assert section.page_height.twips == PAGE_HEIGHT_DXA
    assert section.top_margin.twips == MARGIN_DXA
    assert section.bottom_margin.twips == MARGIN_DXA
    assert section.left_margin.twips == MARGIN_DXA
    assert section.right_margin.twips == MARGIN_DXA
    assert PRESET_NAME == "compact_reference_guide"
    assert HEADER_PATTERN == "editorial_cover"

    title_texts = [paragraph.text for paragraph in document.paragraphs if paragraph.style.name == "Title"]
    assert title_texts == ["离线 SBOM 证据工作台"], "editorial cover must occur exactly once"

    for table in document.tables:
        grid = table._tbl.tblGrid
        widths = [int(node.get(qn("w:w"))) for node in grid.findall(qn("w:gridCol"))]
        assert sum(widths) == CONTENT_WIDTH_DXA
        tblpr = table._tbl.tblPr
        assert tblpr.find(qn("w:tblInd")).get(qn("w:w")) == str(TABLE_INDENT_DXA)
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                tcw = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                assert int(tcw.get(qn("w:w"))) == widths[index]

    for paragraph in document.paragraphs:
        text = paragraph.text.lstrip()
        assert not text.startswith("•"), "fake Unicode bullet found"
        assert not re.match(r"^-\s+", text), "fake hyphen bullet found"


def build() -> Path:
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    document = Document()
    _configure_styles(document)
    _configure_sections(document)
    numbering = _configure_numbering(document)
    _set_document_metadata(document)

    _add_editorial_cover(document)
    _add_markdown_body(document, markdown, numbering)
    _audit_document(document)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    output = build()
    print(output)
